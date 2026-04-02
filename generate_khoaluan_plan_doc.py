from docx import Document
from docx.shared import Pt


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def add_numbered(doc, text):
    doc.add_paragraph(text, style="List Number")


def main():
    doc = Document()
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(12)

    add_heading(
        doc,
        "Huong dan trien khai de tai: DevSecOps + GitOps cho Microservices tren AWS EKS",
        level=1,
    )
    doc.add_paragraph(
        "Tai lieu tong hop ke hoach thuc thi de tai theo timeline KLTN, dap ung yeu cau "
        "microservices (toi thieu 5 service) va bao mat ngay tu buoc tao ha tang."
    )

    add_heading(doc, "1. Kien truc de xuat (dam bao >= 5 backend services)", level=2)
    add_bullet(doc, "Frontend: frontend-web (React + TypeScript + Vite)")
    add_bullet(doc, "API Gateway: api-gateway (Node.js NestJS/Fastify)")
    add_bullet(doc, "auth-service: dang nhap, JWT/OAuth2, RBAC")
    add_bullet(doc, "user-service: ho so nguoi dung va role metadata")
    add_bullet(doc, "product-service: nghiep vu san pham")
    add_bullet(doc, "order-service: nghiep vu don hang")
    add_bullet(doc, "notification-service: email/webhook/event thong bao")
    add_bullet(doc, "Database: MongoDB (uu tien schema/DB rieng theo service)")
    add_bullet(doc, "Cache/session/rate limit: Redis (khuyen nghi)")
    doc.add_paragraph(
        "Tong ket: co 6 backend services (gateway + 5 service nghiep vu), vuot muc toi thieu 5."
    )

    add_heading(doc, "2. Stack cong nghe khuyen nghi", level=2)
    add_bullet(doc, "Frontend: React, TailwindCSS, shadcn/ui, TanStack Query")
    add_bullet(doc, "Backend: Node.js NestJS + Fastify adapter")
    add_bullet(doc, "Container: Docker multi-stage")
    add_bullet(doc, "Orchestration: Kubernetes tren AWS EKS")
    add_bullet(doc, "CI: GitHub Actions (goi nhe, de van hanh)")
    add_bullet(doc, "CD GitOps: ArgoCD")
    add_bullet(doc, "Code quality: SonarQube")
    add_bullet(doc, "Security scan: Trivy + Gitleaks + Checkov/tfsec")
    add_bullet(doc, "Observability: Prometheus + Grafana + Loki + Promtail")

    add_heading(doc, "3. Bao mat ngay tu buoc tao ha tang (IaC Security First)", level=2)
    doc.add_paragraph("Pipeline gate cho Terraform (bat buoc truoc apply):")
    add_numbered(doc, "terraform fmt -check")
    add_numbered(doc, "terraform validate")
    add_numbered(doc, "tflint")
    add_numbered(doc, "tfsec")
    add_numbered(doc, "checkov")
    add_numbered(doc, "Fail pipeline neu ton tai loi High/Critical")

    doc.add_paragraph("Hardening can co trong ha tang:")
    add_bullet(doc, "EKS endpoint private hoac public CIDR rat chat")
    add_bullet(doc, "Node group khong public IP")
    add_bullet(doc, "Ma hoa EBS/S3/Secret bang KMS")
    add_bullet(doc, "Dung IRSA, khong hardcode access key")
    add_bullet(doc, "Security Group theo nguyen tac least privilege")
    add_bullet(doc, "Bat CloudTrail + GuardDuty (muc co ban cho demo)")
    add_bullet(doc, "Khong luu secret trong Git; dung Secrets Manager")

    add_heading(doc, "4. Timeline thuc hien bam de cuong", level=2)
    doc.add_paragraph("Thang 3 - Foundation")
    add_bullet(doc, "Chot domain, API contract, va kien truc microservices")
    add_bullet(doc, "Dung skeleton cho frontend + 6 backend services")
    add_bullet(doc, "Khoi tao Terraform cho VPC, EKS, IAM, ECR")
    add_bullet(doc, "Tao CI co ban cho test/build")

    doc.add_paragraph("Thang 4 - Hoan thien ung dung web truoc")
    add_bullet(doc, "Frontend hien dai: dashboard, auth, quan ly data")
    add_bullet(doc, "Hoan thien logic cho cac service backend")
    add_bullet(doc, "Toi uu MongoDB: index, pagination, projection")
    add_bullet(doc, "Them unit test va integration test co ban")

    doc.add_paragraph("Thang 5 - Day du DevSecOps + GitOps + Monitoring")
    add_bullet(doc, "CI day du: test -> sonar -> scan -> build image -> trivy")
    add_bullet(doc, "IaC security gate truoc terraform apply")
    add_bullet(doc, "ArgoCD dong bo tu gitops repo len EKS")
    add_bullet(doc, "Tich hop Prometheus, Grafana, Loki")

    doc.add_paragraph("Thang 6 - Danh gia, bao cao, bao ve")
    add_bullet(doc, "Do luong chi so va danh gia ket qua")
    add_bullet(doc, "Hoan thien bao cao + slide + demo script")
    add_bullet(doc, "Chuan bi cau hoi phan bien va phuong an tra loi")

    add_heading(doc, "5. Thiet ke frontend hien dai, de tuong tac >5 services", level=2)
    add_bullet(doc, "Trang Login/Register, Dashboard tong quan")
    add_bullet(doc, "Man User Management (user-service)")
    add_bullet(doc, "Man Product/Catalog (product-service)")
    add_bullet(doc, "Man Order Workflow (order-service)")
    add_bullet(doc, "Notification Center (notification-service)")
    add_bullet(doc, "Trang Security/Audit read-only")
    add_bullet(doc, "UI/UX: responsive, dark mode, skeleton loading, toast")

    add_heading(doc, "6. Cau truc repository de trien khai gon gang", level=2)
    add_bullet(doc, "app-repo: frontend + services + shared + workflows")
    add_bullet(doc, "infra-repo: terraform modules + env dev/staging")
    add_bullet(doc, "gitops-repo: manifests K8s + overlays + Argo apps")

    add_heading(doc, "7. KPI danh gia de tai", level=2)
    add_bullet(doc, "He thong co >=5 backend services doc lap")
    add_bullet(doc, "100% CD thong qua GitOps (khong deploy tay)")
    add_bullet(doc, "Moi PR phai qua test + quality + security scan")
    add_bullet(doc, "Muc High/Critical = 0 truoc khi deploy")
    add_bullet(doc, "Co dashboard theo doi latency, error rate, cpu/memory")
    add_bullet(doc, "Co rollback duoc bang ArgoCD")

    add_heading(doc, "8. Checklist hanh dong ngay (tuan dau)", level=2)
    add_numbered(doc, "Chot ten va pham vi 6 services")
    add_numbered(doc, "Tao skeleton frontend + backend")
    add_numbered(doc, "Chay local bang Docker Compose")
    add_numbered(doc, "Khoi tao Terraform ECR/VPC/EKS dev")
    add_numbered(doc, "Viet CI dau tien cho 1 service")
    add_numbered(doc, "Bo sung security scan: gitleaks, tfsec/checkov, trivy")

    add_heading(doc, "9. Luu y ve pham vi va gioi han de tai", level=2)
    add_bullet(doc, "Quy mo demo, chua nham den production lon")
    add_bullet(doc, "Chua toi uu chi phi AWS o muc sau")
    add_bullet(doc, "Chua thuc hien pentest toan dien")
    add_bullet(doc, "Tap trung tinh tu dong, nhat quan, an toan va quan sat he thong")

    output = "Ke_hoach_DevSecOps_GitOps_Microservices_AWS_EKS.docx"
    doc.save(output)
    print(f"Da tao file Word: {output}")


if __name__ == "__main__":
    main()
